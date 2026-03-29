from docx import Document
import datetime


class DocumentGenerator:
    def __init__(self, template_path):
        """
        初始化文档生成器
        :param template_path: 模板文件路径
        """
        self.template_path = template_path
    
    def generate(self, data, output_path):
        """
        生成文档
        :param data: 替换数据
        :param output_path: 输出路径
        """
        doc = Document(self.template_path)

        def replace_text_in_paragraph(paragraph, replacements):
            for key, value in replacements.items():
                for run in paragraph.runs:  # 遍历段落中的每个 run
                    if key in run.text:
                        run.text = run.text.replace(key, value)  # 替换文本

        def replace_text_in_table(table, replacements):
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                    for nested_table in cell.tables:
                        replace_text_in_table(nested_table, replacements)  # 递归处理嵌套表格

        # 替换段落中的文本
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, data)

        # 替换表格中的文本
        for table in doc.tables:
            replace_text_in_table(table, data)

        doc.save(output_path)
